import io
import re

from docx import Document

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _parse_inline_runs(text: str) -> list[tuple[str, bool]]:
    """Split a line on **bold** markers into (text, is_bold) runs."""
    runs = []
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            runs.append((text[pos : m.start()], False))
        runs.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False))
    return runs or [(text, False)]


def _add_runs(paragraph, runs: list[tuple[str, bool]]) -> None:
    for text, bold in runs:
        run = paragraph.add_run(text)
        run.bold = bold


def _format_time(seconds: float) -> str:
    total = int(seconds)
    h, rem = divmod(total, 3600)
    m, s = divmod(rem, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def build_verbatim_docx(transcript: dict) -> io.BytesIO:
    """Render a transcript's segments as a VERBATIM docx (spec 4.1): no LLM
    processing, just the segments in order as `[HH:MM:SS] Speaker N：text`."""
    document = Document()
    document.add_heading(transcript["audio_filename"], level=1)

    for seg in transcript["segments"]:
        ts = _format_time(seg["start"])
        speaker_label = f"{seg['speaker']}：" if seg.get("speaker") else ""
        document.add_paragraph(f"[{ts}] {speaker_label}{seg['text']}")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_output_docx(title: str, content: str) -> io.BytesIO:
    """Render an LLM-generated ACTION_ITEMS/SUMMARY output as a docx.

    SUMMARY content uses a constrained Markdown subset (see SUMMARY_PROMPT in
    app/llm.py): `#`/`##` headings, `-`/`  -` bullets (two levels each), and
    `**bold**` spans. ACTION_ITEMS content doesn't use this syntax, so its
    lines just fall through to plain paragraphs, same as before.
    """
    document = Document()
    document.add_heading(title, level=1)

    for raw_line in content.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            continue

        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=3)
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=2)
            continue

        stripped = line.lstrip(" ")
        indent = len(line) - len(stripped)
        if stripped.startswith("- "):
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            paragraph = document.add_paragraph(style=style)
            _add_runs(paragraph, _parse_inline_runs(stripped[2:].strip()))
            continue

        paragraph = document.add_paragraph()
        _add_runs(paragraph, _parse_inline_runs(line))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
